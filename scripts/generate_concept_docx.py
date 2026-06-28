from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ALBA_GREEN = RGBColor(29, 48, 41)
ALBA_CLAY = RGBColor(184, 95, 50)
ALBA_MUTED = RGBColor(90, 99, 95)
ALBA_TEXT = RGBColor(28, 35, 32)


def set_run_font(run, *, name: str = "Arial", size: float | None = None, bold: bool | None = None,
                 italic: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_paragraph_bottom_border(paragraph, color: str = "D9CEC0", size: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_document(document: Document, footer_label: str) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_run = footer_p.add_run(footer_label)
    set_run_font(footer_run, size=9.5, color=ALBA_MUTED)


def add_title_block(document: Document, title: str, meta_lines: list[str]) -> None:
    title_p = document.add_paragraph()
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run(title)
    set_run_font(title_run, size=24, bold=True, color=ALBA_GREEN)

    for meta in meta_lines:
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(meta)
        set_run_font(run, size=10, color=ALBA_MUTED)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(12)
    add_paragraph_bottom_border(spacer)


def add_heading(document: Document, text: str, level: int) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 2 else 12)
    p.paragraph_format.space_after = Pt(6 if level == 2 else 4)
    run = p.add_run(text)
    if level == 2:
        set_run_font(run, size=16, bold=True, color=ALBA_GREEN)
    elif level == 3:
        set_run_font(run, size=13, bold=True, color=ALBA_CLAY)
    else:
        set_run_font(run, size=12, bold=True, color=ALBA_GREEN)


def add_body_paragraph(document: Document, text: str, *, bold: bool = False) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=11, bold=bold, color=ALBA_TEXT)


def add_list_paragraph(document: Document, text: str, style_name: str) -> None:
    p = document.add_paragraph(style=style_name)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=11, color=ALBA_TEXT)


def render_markdown_to_docx(markdown_path: Path, docx_path: Path, footer_label: str) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_document(document, footer_label)

    title = ""
    meta_lines: list[str] = []
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if not title and line.startswith("# "):
            title = line[2:].strip()
            index += 1
            continue
        if title and line and not line.startswith("#"):
            meta_lines.append(line)
            index += 1
            continue
        break

    add_title_block(document, title, meta_lines)

    bullet_re = re.compile(r"^- (.+)$")
    number_re = re.compile(r"^\d+\. (.+)$")

    for raw_line in lines[index:]:
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("## "):
            add_heading(document, stripped[3:].strip(), 2)
            continue
        if stripped.startswith("### "):
            add_heading(document, stripped[4:].strip(), 3)
            continue
        bullet_match = bullet_re.match(stripped)
        if bullet_match:
            add_list_paragraph(document, bullet_match.group(1), "List Bullet")
            continue
        number_match = number_re.match(stripped)
        if number_match:
            add_list_paragraph(document, number_match.group(1), "List Number")
            continue
        if stripped.startswith("> "):
            add_body_paragraph(document, stripped[2:].strip(), bold=True)
            continue
        if stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            add_body_paragraph(document, stripped[2:-2].strip(), bold=True)
            continue
        add_body_paragraph(document, stripped)

    document.save(docx_path)


def main() -> int:
    if len(sys.argv) != 4:
        print("Usage: generate_concept_docx.py <input.md> <output.docx> <footer_label>")
        return 1

    markdown_path = Path(sys.argv[1]).resolve()
    docx_path = Path(sys.argv[2]).resolve()
    footer_label = sys.argv[3]

    render_markdown_to_docx(markdown_path, docx_path, footer_label)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
